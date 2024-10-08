import fitz  # PyMuPDF
import pypandoc
from docx import Document
from docx.shared import Inches

def pdf_to_word(pdf_path, word_path):
    # Create a Word document
    doc = Document()

    # Open the PDF file
    pdf_document = fitz.open(pdf_path)
    
    for page_num in range(len(pdf_document)):
        # Get the page
        page = pdf_document.load_page(page_num)
        
        # Extract text and add to the Word document
        text = page.get_text("text")
        if text:
            doc.add_paragraph(text)
        
        # Extract images
        images = page.get_images(full=True)
        for img in images:
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]

            # Insert image into the Word document
            image_filename = f"image{page_num}_{xref}.png"
            with open(image_filename, "wb") as image_file:
                image_file.write(image_bytes)
            
            # Add image to Word document
            doc.add_picture(image_filename, width=Inches(5))
    
    # Save the Word document
    doc.save(word_path)
    print(f"PDF converted to Word and saved as {word_path}")

if __name__ == "__main__":
    # Example usage
    pdf_file = "path/to/your/file.pdf"
    word_file = "path/to/save/converted_file.docx"
    
    pdf_to_word(pdf_file, word_file)
